import os, time
import pandas as pd
import joblib
import numpy as np
import matplotlib.pyplot as plt
import fasttext
import lightgbm as lgb
import mlflow
import mlflow.keras
from mlflow.keras import MLflowCallback
from docx import Document
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import GridSearchCV, train_test_split
from sklearn.metrics import accuracy_score, f1_score
from lightgbm import early_stopping, log_evaluation
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import LSTM, Dense, Dropout, Bidirectional
from tensorflow.keras.callbacks import EarlyStopping
import tensorflow as tf
from transformers import DistilBertForSequenceClassification, Trainer, TrainingArguments
from utils import mlflow_run_safety, suivi_temps_ressources




# Logistic Regression
@mlflow_run_safety(experiment_name = "P7_sentiment_analysis")
def train_logistic_regression_with_cv(X, y, model_path="../models_saved/log_reg_model.pkl", force_retrain = False):
    if os.path.exists(model_path) and not force_retrain:
        print("✅ Modèle Régression Logistique déjà existant. Chargement...")
        return joblib.load(model_path)

    start = time.time()
    param_grid = {
        "C": [0.01, 0.1, 1, 10, 100],
        "penalty": ["l1", "l2"],
        "solver": ["liblinear"]
    }
    log_reg = LogisticRegression(max_iter = 200)
    grid_search = GridSearchCV(log_reg, param_grid, cv = 5, scoring = "accuracy", verbose = 1, n_jobs = -1)

    print("🔄 Entraînement Régression Logistique...")
    grid_search.fit(X, y)
    best_log_reg = grid_search.best_estimator_

    # Log hyperparams
    mlflow.log_params(grid_search.best_params_)

    # Évaluation
    y_pred = best_log_reg.predict(X)
    acc = accuracy_score(y, y_pred)
    f1 = f1_score(y, y_pred)
    mlflow.log_metric("final_accuracy", acc)
    mlflow.log_metric("final_f1_score", f1)

    # Sauvegarde locale
    joblib.dump(best_log_reg, model_path)
    print(f"✅ Modèle sauvegardé sous {model_path}")

    # Log modèle dans MLFlow + Registry
    mlflow.sklearn.log_model(best_log_reg, artifact_path = "logreg_model")
    run_id = mlflow.active_run().info.run_id
    model_uri = f"runs:/{run_id}/logreg_model"
    mlflow.register_model(model_uri, "sentiment_model_logreg")

    suivi_temps_ressources(start, "Régression Logistique")
    return best_log_reg


# Random Forest
@mlflow_run_safety(experiment_name = "P7_sentiment_analysis")
def train_random_forest(X_train, y_train, model_path = "../models_saved/rf_model.pkl", force_retrain = False):
    if os.path.exists(model_path) and not force_retrain:
        print("✅ Modèle RandomForest déjà existant. Chargement...")
        return joblib.load(model_path)

    start = time.time()
    rf = RandomForestClassifier(n_estimators = 100, max_depth = 10, n_jobs = -1)
    rf.fit(X_train, y_train)

    # Logging hyperparams
    mlflow.log_params({
        "n_estimators": 100,
        "max_depth": 10
    })

    # Logging métriques
    y_pred = rf.predict(X_train)
    acc = accuracy_score(y_train, y_pred)
    f1 = f1_score(y_train, y_pred)
    mlflow.log_metric("final_accuracy", acc)
    mlflow.log_metric("final_f1_score", f1)

    # Sauvegarde locale
    joblib.dump(rf, model_path)
    print(f"✅ Modèle RandomForest sauvegardé sous {model_path}")

    # Log modèle dans MLFlow + Registry
    mlflow.sklearn.log_model(rf, artifact_path="rf_model")
    run_id = mlflow.active_run().info.run_id
    model_uri = f"runs:/{run_id}/rf_model"
    mlflow.register_model(model_uri, "sentiment_model_rf")

    suivi_temps_ressources(start, "RandomForest")
    return rf


# LightGBM
@mlflow_run_safety(experiment_name = "P7_sentiment_analysis")
def train_lightgbm(X_train, y_train, X_val, y_val, model_path = "../models_saved/lgbm_model.txt", force_retrain = False):
    if os.path.exists(model_path) and not force_retrain:
        print("✅ Modèle LightGBM existant. Chargement...")
        return lgb.Booster(model_file = model_path)

    start = time.time()

    lgb_train = lgb.Dataset(X_train, label = y_train)
    lgb_eval = lgb.Dataset(X_val, label = y_val, reference = lgb_train)

    params = {
        'objective': 'binary',
        'metric': 'binary_logloss',
        'num_leaves': 31,
        'learning_rate': 0.1,
        'max_depth': 5,
        'n_jobs': -1,
        'verbose': -1
    }

    print("🚀 Entraînement LightGBM...")

    lgbm_model = lgb.train(
        params,
        lgb_train,
        valid_sets = [lgb_eval],
        num_boost_round = 100,
        callbacks = [early_stopping(stopping_rounds = 10), log_evaluation(period = 10)]
    )

    mlflow.log_params(params)

    y_val_pred = (lgbm_model.predict(X_val) > 0.5).astype(int)
    acc = accuracy_score(y_val, y_val_pred)
    f1 = f1_score(y_val, y_val_pred)

    mlflow.log_metric("final_accuracy", acc)
    mlflow.log_metric("final_f1_score", f1)

    lgbm_model.save_model(model_path)
    print(f"✅ Modèle LightGBM sauvegardé sous {model_path}")

    # Log modèle dans MLFlow + Registry
    mlflow.lightgbm.log_model(lgbm_model, artifact_path="lgbm_model")
    run_id = mlflow.active_run().info.run_id
    model_uri = f"runs:/{run_id}/lgbm_model"
    mlflow.register_model(model_uri, "sentiment_model_lgbm")

    suivi_temps_ressources(start, "LightGBM")
    return lgbm_model


# FastText Supervised Training
@mlflow_run_safety(experiment_name = "P7_sentiment_analysis")
def train_fasttext_supervised(file_path = "../models_saved/tweets_fasttext.txt", model_path = "../models_saved/fasttext_model.ftz", force_retrain = False):
    if os.path.exists(model_path) and not force_retrain:
        print("✅ Modèle FastText supervisé existant. Chargement...")
        return fasttext.load_model(model_path)

    print("🚀 Entraînement FastText supervisé...")
    model = fasttext.train_supervised(file_path, epoch = 10, wordNgrams = 2, dim = 300)
    model.save_model(model_path)
    print(f"✅ Modèle FastText sauvegardé sous {model_path}")

    # Logging dans MLFlow
    mlflow.log_param("epoch", 10)
    mlflow.log_param("wordNgrams", 2)
    mlflow.log_param("dim", 300)

    mlflow.log_artifact(model_path, artifact_path="fasttext_model")

    run_id = mlflow.active_run().info.run_id
    model_uri = f"runs:/{run_id}/fasttext_model/{os.path.basename(model_path)}"
    mlflow.register_model(model_uri, "sentiment_model_fasttext")

    return model

@mlflow_run_safety(experiment_name="P7_sentiment_analysis")
def train_lstm_model(
    X_embeddings,
    y_labels,
    model_base_path="../models_saved/lstm_model",
    force_retrain=False,
    lstm_units=128,
    dropout_rate=0.3,
    dense_units=64,
    batch_size=256,
    epochs=10,
    optimizer='adam',
    loss_function='binary_crossentropy',
    results_csv_path="../results/lstm_experiments_results.csv",
    resume_doc_path="../results/resume_lstm.docx"
):
    # ✅ Créer les dossiers nécessaires si non existants
    os.makedirs(os.path.dirname(results_csv_path), exist_ok=True)
    os.makedirs(os.path.dirname(resume_doc_path), exist_ok=True)
    
    X = np.array(X_embeddings)
    y = np.array(y_labels).astype(int)

    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, stratify=y, random_state=70)

    X_train_reshaped = X_train.reshape((X_train.shape[0], 1, X.shape[1]))
    X_test_reshaped = X_test.reshape((X_test.shape[0], 1, X.shape[1]))

    # 📁 Génération du nom de fichier basé sur les params
    filename = f"LSTM_{lstm_units}u_{batch_size}bs_{epochs}ep.h5"
    model_path = os.path.join(model_base_path + f"_{lstm_units}_{batch_size}_{epochs}.h5")

    if force_retrain and os.path.exists(model_path):
        os.remove(model_path)

    if os.path.exists(model_path) and not force_retrain:
        print(f"✅ Modèle déjà existant. Chargement : {model_path}")
        model = tf.keras.models.load_model(model_path)
        return model, (X_test_reshaped, y_test), None

    start = time.time()
    mlflow.set_tag("mlflow.runName", f"LSTM_{lstm_units}u_{batch_size}bs_{epochs}ep")

    # Logging des hyperparams
    mlflow.log_param("lstm_units", lstm_units)
    mlflow.log_param("dropout_rate", dropout_rate)
    mlflow.log_param("dense_units", dense_units)
    mlflow.log_param("batch_size", batch_size)
    mlflow.log_param("epochs", epochs)
    mlflow.log_param("optimizer", optimizer)
    mlflow.log_param("loss_function", loss_function)

    model = Sequential()
    model.add(Bidirectional(LSTM(lstm_units), input_shape=(1, X.shape[1])))
    model.add(Dropout(dropout_rate))
    model.add(Dense(dense_units, activation='relu'))
    model.add(Dense(1, activation='sigmoid'))
    model.compile(optimizer=optimizer, loss=loss_function, metrics=['accuracy'])

    early_stop = EarlyStopping(monitor='val_loss', patience=3, restore_best_weights=True)

    print("🚀 Entraînement en cours...")
    history = model.fit(
        X_train_reshaped, y_train,
        validation_data=(X_test_reshaped, y_test),
        epochs=epochs,
        batch_size=batch_size,
        callbacks=[early_stop],
        verbose=1
    )

    # 📊 Log des courbes dans MLflow
    for epoch in range(len(history.history['loss'])):
        mlflow.log_metric("train_loss", history.history['loss'][epoch], step=epoch)
        mlflow.log_metric("val_loss", history.history['val_loss'][epoch], step=epoch)
        mlflow.log_metric("train_accuracy", history.history['accuracy'][epoch], step=epoch)
        mlflow.log_metric("val_accuracy", history.history['val_accuracy'][epoch], step=epoch)

    # 💾 Sauvegarde modèle + log
    model.save(model_path)
    mlflow.keras.log_model(model, artifact_path="lstm_model")
    run_id = mlflow.active_run().info.run_id
    mlflow.register_model(f"runs:/{run_id}/lstm_model", "sentiment_model_lstm")

    # 📈 Écart accuracy train/val (visuel uniquement)
    gap_acc = [tr - va for tr, va in zip(history.history['accuracy'], history.history['val_accuracy'])]
    plt.figure(figsize=(8, 3))
    plt.plot(gap_acc)
    plt.title("Écart accuracy (train - val)")
    plt.xlabel("Epoch")
    plt.grid(True)
    plt.tight_layout()
    gap_plot_path = model_path.replace(".h5", "_gap_acc.png")
    plt.savefig(gap_plot_path)
    mlflow.log_artifact(gap_plot_path)

    # 📊 Mise à jour CSV des résultats
    y_pred = (model.predict(X_test_reshaped) > 0.5).astype(int).flatten()
    acc = round(accuracy_score(y_test, y_pred), 4)
    f1 = round(f1_score(y_test, y_pred), 4)

    new_row = pd.DataFrame([{
        "model_file": os.path.basename(model_path),
        "lstm_units": lstm_units,
        "dropout_rate": dropout_rate,
        "dense_units": dense_units,
        "batch_size": batch_size,
        "epochs": epochs,
        "accuracy": acc,
        "f1_score": f1
    }])

    if os.path.exists(results_csv_path):
        df = pd.read_csv(results_csv_path)
        df = pd.concat([df, new_row], ignore_index=True)
    else:
        df = new_row
    df.to_csv(results_csv_path, index=False)

    # 🥇 Génération résumé Word avec la meilleure config
    best = df.sort_values("f1_score", ascending=False).iloc[0]
    doc = Document()
    doc.add_heading("Résumé des Entraînements LSTM", level=1)
    doc.add_heading("Meilleure configuration", level=2)
    for col in best.index:
        doc.add_paragraph(f"{col} : {best[col]}")
    doc.save(resume_doc_path)

    # 💾 Sauvegarde de la meilleure config LSTM
    joblib.dump(best.to_dict(), '../models_saved/best_lstm_config.pkl')
    print("✅ Configuration LSTM optimale sauvegardée dans best_lstm_config.pkl")

    # 🕒 Temps total
    suivi_temps_ressources(start, "LSTM")

    return model, (X_test_reshaped, y_test), history



# LSTM Training sur FastText



# @mlflow_run_safety(experiment_name="P7_sentiment_analysis")
# def train_lstm_model(X_embeddings, y_labels, model_path = "../models_saved/lstm_model.h5", force_retrain = False):
#     X = np.array(X_embeddings)
#     y = np.array(y_labels).astype(int)

#     X_train, X_test, y_train, y_test = train_test_split(X, y, test_size = 0.2, stratify = y, random_state = 70)

#     X_train_reshaped = X_train.reshape((X_train.shape[0], 1, X_train.shape[1]))
#     X_test_reshaped = X_test.reshape((X_test.shape[0], 1, X_test.shape[1]))

#     if os.path.exists(model_path) and not force_retrain:
#         print(f"✅ Modèle LSTM déjà existant. Chargement...")
#         model = tf.keras.models.load_model(model_path)
#         return model, (X_test_reshaped, y_test), None

#     start = time.time()
#     model = Sequential()
#     model.add(Bidirectional(LSTM(128), input_shape=(1, X.shape[1])))
#     model.add(Dropout(0.3))
#     model.add(Dense(64, activation = 'relu'))
#     model.add(Dense(1, activation = 'sigmoid'))

#     model.compile(optimizer = 'adam', loss = 'binary_crossentropy', metrics = ['accuracy'])
#     early_stop = EarlyStopping(monitor = 'val_loss', patience = 3, restore_best_weights = True)
#     mlflow_callback = MLflowCallback()

#     print("🚀 Entraînement LSTM...")
#     history = model.fit(
#         X_train_reshaped, y_train,
#         validation_data = (X_test_reshaped, y_test),
#         epochs = 10, batch_size = 256,
#         callbacks = [early_stop, mlflow_callback],
#         verbose = 1
#     )

#     model.save(model_path)
#     print(f"✅ Modèle LSTM sauvegardé sous {model_path}")

#     # Logging du modèle
#     mlflow.keras.log_model(model, artifact_path="lstm_model")
#     run_id = mlflow.active_run().info.run_id
#     model_uri = f"runs:/{run_id}/lstm_model"
#     mlflow.register_model(model_uri, "sentiment_model_lstm")

#     suivi_temps_ressources(start, "LSTM")
#     return model, (X_test_reshaped, y_test), history


# DistilBERT Fine-tuning
@mlflow_run_safety(experiment_name="P7_sentiment_analysis")
def train_distilbert_model(tokenized_dataset, model_save_path = "../models_saved/distilbert_model", force_retrain=False):
    from datasets import ClassLabel
    if os.path.exists(model_save_path) and not force_retrain:
        print(f"✅ Modèle DistilBERT déjà fine-tuné. Chargement depuis {model_save_path}...")
        model = DistilBertForSequenceClassification.from_pretrained(model_save_path)
        return model, None, None

    print("🚀 Fine-tuning DistilBERT...")
    model = DistilBertForSequenceClassification.from_pretrained("distilbert-base-uncased", num_labels = 2)

    if tokenized_dataset.features['label'].__class__.__name__ != 'ClassLabel':
        num_classes = len(set(tokenized_dataset['label']))
        tokenized_dataset = tokenized_dataset.cast_column('label', ClassLabel(num_classes = num_classes))

    dataset_split = tokenized_dataset.train_test_split(test_size = 0.2, stratify_by_column = 'label')
    train_dataset = dataset_split['train']
    test_dataset = dataset_split['test']

    training_args = TrainingArguments(
        output_dir = "../models_saved/distilbert_output",
        evaluation_strategy = "epoch",
        save_strategy = "epoch",
        learning_rate = 2e-5,
        per_device_train_batch_size = 16,
        per_device_eval_batch_size = 16,
        num_train_epochs = 3,
        weight_decay = 0.01,
        save_total_limit = 1,
        load_best_model_at_end = True,
        metric_for_best_model = "accuracy",
        logging_dir = "../models_saved/logs",
        logging_steps = 50,
        report_to = "mlflow"
    )

    def compute_metrics(eval_pred):
        logits, labels = eval_pred
        predictions = logits.argmax(axis = -1)
        acc = accuracy_score(labels, predictions)
        f1 = f1_score(labels, predictions)
        return {"accuracy": acc, "f1": f1}

    trainer = Trainer(
        model = model,
        args = training_args,
        train_dataset = train_dataset,
        eval_dataset = test_dataset,
        compute_metrics = compute_metrics
    )

    trainer.train()
    model.save_pretrained(model_save_path)
    print(f"✅ Modèle DistilBERT sauvegardé sous {model_save_path}")

    mlflow.pytorch.log_model(model, artifact_path="distilbert_model")
    run_id = mlflow.active_run().info.run_id
    model_uri = f"runs:/{run_id}/distilbert_model"
    mlflow.register_model(model_uri, "sentiment_model_distilbert")

    return model, trainer, test_dataset